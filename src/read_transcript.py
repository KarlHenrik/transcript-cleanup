from pathlib import Path
from docx import Document
from src.create_docx import get_cells_grid
import re
from itertools import repeat
import docx

def read_transcript(filename):
    p = Path(filename)
    if p.suffix == ".vtt":
        return read_vtt(filename)
    elif p.suffix == ".docx":
        return decide_docx_format(filename)
    else:
        raise ValueError("Invalid file extenstion. Only accepts .vtt or .docx files")

def decide_docx_format(filename):
    document = Document(filename)

    if len(document.tables) == 0:
        fullText = []
        for para in document.paragraphs:
            fullText.append(para.text)
        raw_text = '\n'.join(fullText)
        if raw_text[0] == "[":
            return read_MAXQDA_Text(filename, raw_text)
        else:
            return read_Timed_Text(filename, raw_text)

    table = document.tables[0]
    try:
        cells = get_cells_grid(table)
    except docx.oxml.exceptions.InvalidXmlError:
        raise NotImplementedError(f"You need to open and save the file {filename} in Word first, sorry!")

    headers = [cell.text for cell in cells[0]]
    cells = cells[1:]
    match headers:
        case ["Time", "", "ID", "Content"]:
            return read_Timed(filename, cells)
        case ["", "ID", "Content"]:
            if has_MAXQDA_timestamp(cells[0][-1].text): #MAXQDA time stamp in first cell
                return read_MAXQDA(filename, cells)
            elif has_MAXQDA_timestamp(cells[1][-1].text): #Mangled Timestamps due to MAXQDA export
                first_time = document.paragraphs[0].text.strip()
                return read_MAXQDA_export(filename, cells, first_time)
            else:
                return read_NoTime(filename, cells)
        
    raise ValueError("File has invalid formatting")

def time_from_vtt(line_with_time):
    return line_with_time.strip().split(" --> ")[0][:-2] # hh:mm:ss:x

def read_vtt(filename):
    """
    Reading the vtt file and turning it into lists of strings
    """
    with open(filename, "r") as ifile:
        ifile.readline() # WEBVTT
        ifile.readline() # BLANK
        text = ifile.readlines()

    contents = [text[1].strip()]
    times = [time_from_vtt(text[0])]

    sentence_completed = contents[0][-1] in [".","?","!"]
    for i in range(3, len(text) - 1, 3):
        started_new_sentence = False
        content = text[i+1].strip()
        new_contents = [c for c in re.finditer('\.\.\.|\.|\?|\!', content)]

        if len(new_contents) == 0: # No punctuation means incomplete sentence
            if sentence_completed: # Starts new sentence without end
                contents.append(content)
                times.append(time_from_vtt(text[i]))
            else:
                contents[-1] = contents[-1] + " " + content.strip() # Just keeping sentence going
            sentence_completed = False
            continue

        start = 0
        first_new_sentence = 0
        #If the first new content just there to complete the sentence?
        if not sentence_completed:
            contents[-1] = contents[-1] + " " + content[start:new_contents[0].end(0)].strip()
            start = new_contents[0].end(0)
            first_new_sentence = 1
        sentence_completed = True

        # Add all new complete sentences
        for c in new_contents[first_new_sentence:]:
            contents.append(content[start:c.end(0)].strip())
            start = c.end(0)
            if started_new_sentence:
                times.append("")
            else:
                times.append(time_from_vtt(text[i]))
            started_new_sentence = True
        end = len(content)

        #If there is more left at the end, it becomes new incomplete sentence
        if len(content[start:end].strip()) > 0: 
            contents.append(content[start:end].strip())
            if started_new_sentence:
                times.append("")
            else:
                times.append(time_from_vtt(text[i]))
            sentence_completed = False

    times = [f"00:{time}" if len(time) == 7 else time for time in times]
    IDs = [i for i in repeat("", len(contents))]
    return contents, IDs, times

def read_Timed_Text(filename, s):
    print(f"Reading {filename} as Timed Text format...")
    times = []
    contents = []
    IDs = []
    bad_id = re.compile(r"\n\d\d", re.IGNORECASE)
    pattern = re.compile(r"\d\d:\d\d:\d\d\.\d[^:]+:", re.IGNORECASE)

    time_ID_locs = [(m.start(0), m.end(0)) for m in re.finditer(pattern, s)]
    ends = [m[0] for m in time_ID_locs[1:]] + [-1]
    for m, end in zip(time_ID_locs, ends):
        times.append(s[m[0]:m[0]+10])
        maybeID = s[m[0]+10:m[1]-1].strip()
        if bad_id.match(maybeID[-3:]):
            IDs.append("")
            contents.append(s[m[0]+10:m[1]-3].strip())
        else:
            if maybeID == "Interviewee":
                maybeID = "I"
            elif maybeID == "Researcher":
                maybeID = "R"
            IDs.append(maybeID)
            contents.append(s[m[1]:end].strip())

    return contents, IDs, times

def read_MAXQDA_Text(filename, s):
    def replacement(match):
        return f"0{match.group(0)[1:-2]}\n"

    mq = re.compile(r"\[\d:\d\d:\d\d\.\d] ", re.IGNORECASE)
    s = mq.sub(replacement, s)

    return read_Timed_Text(filename, s)
    
def read_Timed(filename, cells):
    print(f"Reading {filename} as TimeColumn format...")
    times = [row[0].text for row in cells]
    IDs = [row[2].text for row in cells]
    contents = [row[3].text for row in cells]

    return contents, IDs, times

def read_NoTime(filename, cells):
    print(f"Reading {filename} as NoTime format...")
    IDs = [row[1].text for row in cells]
    contents = [row[2].text for row in cells]
    times = None
    return contents, IDs, times

def has_MAXQDA_timestamp(input_text):
    pattern = re.compile(r".*:.*:.*\..*", re.IGNORECASE)
    return bool(pattern.match(input_text))

def read_MAXQDA(filename, cells):
    print(f"Reading {filename} as MAXQDA format...")
    contents = []
    IDs = [row[1].text for row in cells]
    times = []

    for row in cells:
        raw_content = row[2].text
        timestop = raw_content.find(".") + 2
        times.append(raw_content[:timestop])
        contents.append(raw_content[timestop:])

    return contents, IDs, times

def read_MAXQDA_export(filename, cells, first_time):
    print(f"Reading {filename} as MAXQDA timestamp exported format...")
    contents = [cells[0][2].text]
    IDs = [row[1].text for row in cells]
    times = [first_time]
    
    for row in cells[1:]:
        raw_content = row[2].text
        timestop = raw_content.find(".") + 2
        times.append(raw_content[1:timestop])
        contents.append(raw_content[timestop+1:])

    return contents, IDs, times