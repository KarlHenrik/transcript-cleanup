"""
Creating the word document with the table
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.table import _Cell
from docx.enum.table import WD_TABLE_ALIGNMENT


def styled_doc():
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    return document

def styled_table(document, num_rows, num_cols):
    table = document.add_table(rows=num_rows + 1, cols=num_cols)
    table.autofit = False 
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #table.style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    table.style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    table.style.paragraph_format.line_spacing = 1.15

    return table

def set_table_widths(cells, widths):
    for row in cells:
        for idx, width in enumerate(widths):
            row[idx].width = width

def get_cells_grid(table):
    cells = [[]]
    col_count = table._column_count
    for tc in table._tbl.iter_tcs():
        cells[-1].append(_Cell(tc, table))
        if len(cells[-1]) == col_count:
            cells.append([])
    return cells[:-1]

def create_docx(filename, headers, rows, widths):
    num_cols = len(headers)
    num_rows = len(rows[0])

    document = styled_doc()
    table = styled_table(document, num_rows, num_cols)

    cells = get_cells_grid(table)
    set_table_widths(cells, widths)
    
    hdr_cells = cells[:][0]
    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr

    for i, row in enumerate(zip(*rows)):
        row_cells = cells[:][i+1]
        for j, r in enumerate(row):
            row_cells[j].text = str(r)

    
    document.save(filename)
import random
def create_docx_text(filename, times, contents, IDs):
    document = styled_doc()
    
    for time, content, ID in zip(times, contents, IDs):
        p = document.add_paragraph(time + "\n")
        if ID:
            p.add_run(ID + ": ").bold = True
        p.add_run(content)

    document.save(filename)