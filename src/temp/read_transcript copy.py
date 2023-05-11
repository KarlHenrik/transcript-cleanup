
from pathlib import Path
from docx import Document
from src.create_docx import get_cells_grid
from src.transcript_table import Timed_Table, NoTime_Table
import re

from enum import Flag

class Format(Flag):
    MAXQDA = 1 # Time | (idx) | ID | Content
    Timed = 2 # (idx) | ID | [Time]Content
    NoTime = 3 # (idx) | ID | Content
    VTT = 4
    MAXQDA_Mangled_Export = 5

InputFormat = 

def read_transcript(filename):
    p = Path(filename)
    if p.suffix == ".vtt":
        return read_vtt(filename)
    elif p.suffix == ".docx":
        return read_docx(filename)
    else:
        raise ValueError("Invalid file extenstion. Only accepts .vtt or .docx files")

def start_time_from_vtt(line_with_time):
    return line_with_time.strip().split(" --> ")[0]

def read_vtt(filename):
    """
    Reading the vtt file and turning it into lists of strings
    """
    with open(filename, "r") as ifile:
        ifile.readline() # WEBVTT
        ifile.readline() # BLANK
        text = ifile.readlines()

    times = [start_time_from_vtt(text[0])]
    contents = [text[1].strip()]

    last_usable_index = len(text) - 1
    for i in range(3, last_usable_index, 3):
        
        content = text[i+1].strip()

        if len(contents[-1]) == 0 or contents[-1][-1] not in [".", "!", "?"]:
            contents[-1] = contents[-1] + " " + content
        else:
            contents.append(content)
            times.append(start_time_from_vtt(text[i]))
        
    return Timed_Table(filename, times, contents)


def find_MAXQDA_timestamp(input_text):
    pattern = re.compile(r"\[.*:.*:.*\..*\]", re.IGNORECASE)
    return bool(pattern.match(input_text))

def read_docx(filename):
    document = Document(filename)
    table = document.tables[0]
    try:
        cells = get_cells_grid(table)
    except docx.oxml.exceptions.InvalidXmlError:
        raise NotImplementedError(f"You need to open and save the file {filename} in Word first, sorry!")

    headers = [cell.text for cell in cells[0]]
    cells = cells[1:]
    match headers:
        case ["Time", "", "ID", "Content"]:
            return read_Timed(filename, cells)
        case ["", "ID", "Content"]:
            if find_MAXQDA_timestamp(cells[0][-1].text):
                return read_MAXQDA(filename, cells)
            elif find_MAXQDA_timestamp(cells[1][-1].text):
                first_time = document.paragraphs[0].text[1:-1]
                return read_MAXQDA_exported(filename, first_time, cells)
            else:
                return read_NoTime(filename, cells)
        
    raise ValueError("File has invalid formatting")
    
def read_Timed(filename, cells):
    print(f"Reading {filename} as TimeColumn format...")
    times = [row[0].text for row in cells]
    IDs = [row[2].text for row in cells]
    contents = [row[3].text for row in cells]

    return Timed_Table(filename, times, contents, IDs)

def read_NoTime(filename, cells):
    print(f"Reading {filename} as NoTime format...")
    IDs = [row[1].text for row in cells]
    contents = [row[2].text for row in cells]

    return NoTime_Table(filename, contents, IDs)

def read_MAXQDA(filename, cells):
    print(f"Reading {filename} as MAXQDA format...")
    IDs = [row[1].text for row in cells]

    times = []
    contents = []
    for row in cells:
        raw_content = row[2].text
        timestop = raw_content.find("]")
        assert timestop != -1
        times.append(raw_content[1:timestop])
        contents.append(raw_content[timestop+1:])

    return Timed_Table(filename, times, contents, IDs)

def read_MAXQDA_exported(filename, first_time, cells):
    print(f"Reading {filename} as MAXQDA export...")
    IDs = [row[1].text for row in cells]

    times = [first_time]
    contents = [cells[0][2].text]
    for row in cells[1:]:
        raw_content = row[2].text
        timestop = raw_content.find("]")
        assert timestop != -1
        times.append(raw_content[1:timestop])
        contents.append(raw_content[timestop+1:])

    return Timed_Table(filename, times, contents, IDs)